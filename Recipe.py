import docx, requests      # These are the libraries that are used in the files to create them in Python


'''Down here, I will create the function that will get the taco API and find the specific parts to 
take to the Word Document that will be created.'''


url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'      ''' This is the URL of the API and 
                                                                        how I will identify it in the code'''

data = requests.get(url).json()    # This is the API where I get the random taco recipes

seasoning_data = data["seasoning"]['recipe']    # Gathering the data for the seasoning
seasoning_name = data["seasoning"]['name']      # This is where I get just the name for the seasoning
condiment_data = data['condiment']['recipe']    # Gathering data for the condiments
condiment_name = data['condiment']['name']      # Getting the name of the condiments
mixin_data = data['mixin']['recipe']            # This is data for the mixing
mixin_name = data['mixin']['name']              # This is just the name for the mixin
base_layer_data = data['base_layer']['recipe']  # This is the data for the base layer
base_name = data['base_layer']['name']          # This is the name for the base layer
shell_data = data['shell']['recipe']            # Gathering data for the shell
shell_name = data['shell']['name']              # This is where I get the name of the shell

'''This upcoming section is where my code creates
    the Word document'''
document = docx.Document()

document.add_paragraph('Random Taco Cookbook', 'Title')     # That is the name of the taco cookbook
document.add_picture('Omega Tacos.jpg')     # This is my edited image I received from Unsplash
document.add_paragraph('Picture by: Jason Leung')   # Citing the photographer of the image
document.add_paragraph('Recipes found at: https://taco-1150.herokuapp.com/random/?full_taco=true')  # Citing the API that creates the recipes
document.add_paragraph('Code created by: Paul Snowdey Jr.')         #Citing myself as creator of the code
document.add_page_break()       # This is where the recipes will start


'''The code below is where I create a loop
    for the taco recipes so the Word
    document creates 3 different
    recipes.'''


for taco in range(3):
    '''I copied the data from the recipe to here to create the loop
        to get three different recipes. Without it, it would only print out the same recipe three times'''

    url = 'https://taco-1150.herokuapp.com/random/?full_taco=true'
    data = requests.get(url).json()
    seasoning_data = data["seasoning"]['recipe']
    seasoning_name = data["seasoning"]['name']
    condiment_data = data['condiment']['recipe']
    condiment_name = data['condiment']['name']
    mixin_data = data['mixin']['recipe']
    mixin_name = data['mixin']['name']
    base_layer_data = data['base_layer']['recipe']
    base_name = data['base_layer']['name']
    shell_data = data['shell']['recipe']
    shell_name = data['shell']['name']
    document.add_paragraph(f'Taco {taco + 1}', 'Heading6')      # I made this heading to identify the different recipes
    document.add_paragraph(seasoning_name, 'Title') # This puts the name of the seasoning in a style so it's easier to distinguish
    document.add_paragraph(seasoning_data)  # This is the actually recipe of the seasoning
    document.add_paragraph(mixin_name, 'Title') # Same with the mixin name, different style to identify
    document.add_paragraph(mixin_data)  # The recipe for the mixin
    document.add_paragraph(base_name, 'Title')  # Base layer name in a different style
    document.add_paragraph(base_layer_data) # Actual base layer data
    document.add_paragraph(shell_name, 'Title') # Shell recipe name in a different style
    document.add_paragraph(shell_data)  # Actual shell data and recipe
    document.add_page_break()   #Page break to start the next recipe

document.save('TacoFinish.docx')    #Save the document for the finish.